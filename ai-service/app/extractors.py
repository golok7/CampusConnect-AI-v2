from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO

from fastapi import HTTPException


@dataclass
class ExtractedText:
    text: str
    warnings: list[str] = field(default_factory=list)


# Tolerance for grouping PDF text blocks into rows (points)
_ROW_TOLERANCE = 5

# If right-column blocks outnumber left by less than this ratio, treat as two-column
_TWO_COL_MIN_RATIO = 0.25


def _extract_page_text(page, text_blocks: list) -> str:
    """
    Column-aware text extraction.
    Two-column resumes (common for student CVs) get interleaved by naive y-sort,
    causing section detection to scramble Education into Experience.
    Fix: detect column split, emit left column fully before right column.
    """
    page_width = page.rect.width
    midpoint = page_width * 0.52  # slight right bias — avoids splitting centred headers

    left = [b for b in text_blocks if (b[0] + b[2]) / 2 < midpoint]
    right = [b for b in text_blocks if (b[0] + b[2]) / 2 >= midpoint]

    # Only treat as two-column if both sides have meaningful content
    total = len(text_blocks)
    is_two_col = (
        total >= 6
        and len(left) / total >= _TWO_COL_MIN_RATIO
        and len(right) / total >= _TWO_COL_MIN_RATIO
    )

    if is_two_col:
        left.sort(key=lambda b: b[1])
        right.sort(key=lambda b: b[1])
        return "\n".join(b[4] for b in left) + "\n" + "\n".join(b[4] for b in right)

    # Single-column: row-then-x order
    text_blocks.sort(key=lambda b: (int(b[1] / _ROW_TOLERANCE), b[0]))
    return "\n".join(b[4] for b in text_blocks)


class PdfExtractor:
    def extract(self, file_bytes: bytes) -> ExtractedText:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError("PyMuPDF (pymupdf) is not installed")

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Failed to open PDF: {exc}")

        pages: list[str] = []
        warnings: list[str] = []

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")  # list of (x0, y0, x1, y1, text, ...)
            text_blocks = [b for b in blocks if b[4].strip()]

            if not text_blocks:
                warnings.append(
                    f"Page {page_num} contains no extractable text (OCR not supported)"
                )
                continue

            pages.append(_extract_page_text(page, text_blocks))

        doc.close()
        return ExtractedText(text="\n".join(pages), warnings=warnings)


class DocxExtractor:
    def extract(self, file_bytes: bytes) -> ExtractedText:
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx is not installed")

        try:
            document = docx.Document(BytesIO(file_bytes))
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Failed to open DOCX: {exc}")

        parts: list[str] = []

        # Body paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table cells
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        return ExtractedText(text="\n".join(parts), warnings=[])


class TextExtractor:
    def __init__(self) -> None:
        self._pdf = PdfExtractor()
        self._docx = DocxExtractor()

    def extract(self, file_bytes: bytes, mime_type: str) -> ExtractedText:
        if mime_type == "application/pdf":
            return self._pdf.extract(file_bytes)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._docx.extract(file_bytes)
        else:
            raise HTTPException(status_code=422, detail=f"Unsupported MIME type: {mime_type}")
